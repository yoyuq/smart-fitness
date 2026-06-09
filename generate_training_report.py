from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "鏅鸿兘鍋ヨ韩鎸囧绯荤粺瀹炶鎶ュ憡_鏈€缁堢増.docx"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = '瀹嬩綋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '瀹嬩綋')
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r = p.add_run(code.strip())
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = '瀹嬩綋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '瀹嬩綋')
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.name = '瀹嬩綋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '瀹嬩綋')
        run.font.size = Pt(10.5)


def set_doc_style(doc):
    styles = doc.styles
    styles['Normal'].font.name = '瀹嬩綋'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '瀹嬩綋')
    styles['Normal'].font.size = Pt(10.5)
    for name in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[name].font.name = '榛戜綋'
        styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), '榛戜綋')
        styles[name].font.color.rgb = RGBColor(0, 0, 0)
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 2'].font.size = Pt(14)
    styles['Heading 3'].font.size = Pt(12)


def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('鏅鸿兘鍋ヨ韩鎸囧绯荤粺')
    r.bold = True
    r.font.name = '榛戜綋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '榛戜綋')
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('瀹炶鎶ュ憡')
    r.bold = True
    r.font.name = '榛戜綋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '榛戜綋')
    r.font.size = Pt(22)

    doc.add_paragraph()
    doc.add_paragraph()

    info = [
        ('椤圭洰鍚嶇О', 'Smart Fitness 鏅鸿兘鍋ヨ韩鎸囧绯荤粺'),
        ('椤圭洰绫诲瀷', 'Android APP + FastAPI 鍚庣 + AI 瑙嗚鎺ㄧ悊 + ESP32-CAM/鎵嬫満/PC 澶氭憚鍍忓ご'),
        ('瀹炶鏃堕棿', '2026 骞?5 鏈?- 2026 骞?6 鏈?),
        ('鎶ュ憡鏃ユ湡', '2026 骞?6 鏈?9 鏃?),
        ('浜や粯鐗堟湰', 'v2.0 澶氭憚鍍忓ご闂幆楠岃瘉鐗?),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        set_cell_text(table.rows[i].cells[0], k, True)
        set_cell_text(table.rows[i].cells[1], v)
        set_cell_shading(table.rows[i].cells[0], 'EAF2F8')
    doc.add_page_break()


def main():
    doc = Document()
    set_doc_style(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    title_page(doc)

    doc.add_heading('鎽樿', level=1)
    doc.add_paragraph(
        '鏈疄璁畬鎴愪簡涓€濂椻€滄櫤鑳藉仴韬寚瀵肩郴缁熲€濈殑璁捐銆佸紑鍙戙€佽仈璋冧笌楠岃瘉銆傜郴缁熶互 FastAPI 鍚庣涓轰腑蹇冿紝'
        '缁撳悎 MediaPipe 濮挎€佷及璁°€丄ndroid 鍘熺敓瀹㈡埛绔€丒SP32-CAM 杈圭紭鎽勫儚澶淬€佹墜鏈烘湰鏈烘憚鍍忓ご涓?PC 鎽勫儚澶寸槮瀹㈡埛绔紝'
        '瀹炵幇浜嗚繍鍔ㄥ浘鍍忛噰闆嗐€佽澶囩鍘嬬缉涓婁紶銆佸悗绔粺涓€鎺ㄧ悊銆佸姩浣滆鏁般€佸Э鎬佽瘎鍒嗐€乄ebSocket 瀹炴椂鎺ㄩ€佷互鍙婄Щ鍔ㄧ HUD 灞曠ず銆?
        '鍦ㄦ渶缁堢増鏈腑锛岀郴缁熶粠鍗曚竴 ESP32-CAM 鎵╁睍涓?ESP32-CAM銆丳hone Camera銆丳C Camera 涓夌鎽勫儚澶存潵婧愶紝'
        '骞朵繚鎸佺粺涓€鐨?device_id銆佽缁冧細璇濆拰鎺ㄧ悊鎺ュ彛锛屼粠鑰屾敮鎸佹洿鐏垫椿鐨勫疄璁紨绀哄満鏅€?
    )
    doc.add_paragraph(
        '缁忛獙璇侊紝Android 妯℃嫙鍣ㄤ腑 Phone Camera 妯″紡鍙畬鎴?CameraX 閲囬泦銆佸帇缂┿€佷笂浼犮€佹帹鐞嗕笌璁粌鎺у埗闂幆锛?
        'PC Camera 妯″紡鍙€氳繃鏈満 OpenCV 閲囬泦鑴氭湰涓婁紶鐢婚潰锛孉PP 绔闃呭悗绔疄鏃剁粨鏋滃苟鏄剧ず璁℃暟銆佽瘎鍒嗗拰鍙嶉锛?
        'ESP32-CAM 鍘熸湁璺緞淇濇寔鍏煎锛岄€傜敤浜庣湡鏈虹‖浠惰仈璋冨拰鐜板満婕旂ず銆?
    )

    doc.add_heading('涓€銆佸疄璁洰鐨勪笌浠诲姟瑕佹眰', level=1)
    add_numbered(doc, [
        '鎺屾彙绉诲姩绔€佸悗绔€佸祵鍏ュ紡璁惧涓?AI 瑙嗚妯″潡鐨勭患鍚堢郴缁熼泦鎴愭柟娉曘€?,
        '瀹炵幇鍩轰簬鎽勫儚澶寸殑杩愬姩濮挎€佽瘑鍒€佸姩浣滆鏁般€佸Э鍔胯瘎鍒嗗拰瀹炴椂鍙嶉銆?,
        '瀹屾垚 Android APP 涓庡悗绔帴鍙ｃ€乄ebSocket 鎺ㄩ€併€佽缁冪姸鎬佹帶鍒朵箣闂寸殑鑱旇皟銆?,
        '瀹炵幇 ESP32-CAM銆佹墜鏈烘憚鍍忓ご銆丳C 鎽勫儚澶翠笁绉嶅浘鍍忔潵婧愮殑缁熶竴鎺ュ叆銆?,
        '淇濊瘉鍥惧儚鍦ㄦ憚鍍忓ご鎵€鍦ㄨ澶囩瀹屾垚 resize 涓?JPEG 鍘嬬缉鍚庡啀涓婁紶锛岄檷浣庣綉缁滃拰鍚庣鍘嬪姏銆?,
        '杈撳嚭鍙繍琛?APK銆佸悗绔湇鍔°€丳C 鎽勫儚澶村鎴风銆丒SP32 鍥轰欢鍜岄」鐩疄璁姤鍛娿€?,
    ])

    doc.add_heading('浜屻€佺郴缁熸€讳綋璁捐', level=1)
    doc.add_heading('2.1 鎬讳綋鏋舵瀯', level=2)
    doc.add_paragraph(
        '绯荤粺閲囩敤鈥滃绔噰闆?+ 鍚庣缁熶竴鎺ㄧ悊 + APP 瀹炴椂灞曠ず鈥濈殑鏋舵瀯銆傚悇鎽勫儚澶存潵婧愯礋璐ｉ噰闆嗕笌鍘嬬缉锛?
        'FastAPI 鍚庣璐熻矗鐢ㄦ埛銆佽澶囥€佽缁冪姸鎬併€佽瑙夋帹鐞嗗拰 WebSocket 骞挎挱锛孉ndroid APP 璐熻矗璁粌浜や簰涓庣粨鏋滃睍绀恒€?
    )
    add_code(doc, r'''
[ESP32-CAM] --MJPEG/HTTP JPEG-->
[Phone Camera] --CameraX JPEG Base64--> [FastAPI Backend] --> [MediaPipe Pose]
[PC Camera] ----OpenCV JPEG Base64---->        |              --> [Rep Counter/Form Score]
                                                v
                                      [WebSocket Coach Channel]
                                                v
                                      [Android APP HUD/Preview]
''')
    add_table(doc, ['灞傛', '缁勬垚', '涓昏鑱岃矗'], [
        ['鎰熺煡灞?, 'ESP32-CAM銆丄ndroid CameraX銆丳C OpenCV 鎽勫儚澶?, '閲囬泦鍥惧儚銆佺渚?resize銆丣PEG 鍘嬬缉銆佹寜 device_id 涓婁紶'],
        ['鏈嶅姟灞?, 'FastAPI銆丼QLite銆乄ebSocket銆丮QTT', '璁よ瘉銆佽澶囩鐞嗐€佽缁冩帶鍒躲€佹帹鐞嗚仛鍚堛€佸疄鏃跺箍鎾?],
        ['AI 灞?, 'MediaPipe Pose銆丒xerciseDetector銆丗ormAnalyzer', '浜轰綋鍏抽敭鐐规彁鍙栥€佸姩浣滆瘑鍒€佽鏁般€佸Э鎬佽瘎鍒嗕笌鍙嶉'],
        ['搴旂敤灞?, 'Android APP', '鐧诲綍銆佽缁冨紑濮?缁撴潫銆佹憚鍍忓ご鏉ユ簮閫夋嫨銆侀鏋堕瑙堛€佽鏁拌瘎鍒嗗睍绀?],
    ], widths=[1.2, 2.2, 3.8])

    doc.add_heading('2.2 鏍稿績鏁版嵁娴?, level=2)
    add_numbered(doc, [
        '鐢ㄦ埛鍦?Android APP 鐨勮缁冮〉閫夋嫨杩愬姩绫诲瀷鍜屾憚鍍忓ご鏉ユ簮銆?,
        'APP 璋冪敤 /api/v2/training/start锛屽悗绔褰曞綋鍓?device_id 鐨?active training銆?,
        '鎽勫儚澶存潵婧愰噰闆嗗浘鍍忓苟鍦ㄧ渚у帇缂╋紝鐒跺悗璋冪敤 /api/v2/vision/infer/full 涓婁紶 Base64 JPEG銆?,
        '鍚庣浣跨敤 MediaPipe 鎻愬彇鍏抽敭鐐癸紝骞舵寜 exercise/device_id 杩涜璁℃暟鍜岃瘎鍒嗐€?,
        '鍚庣灏?rep_count銆乫orm_score銆乫eedback銆乨evice_id銆乻ource 绛夌粨鏋滈€氳繃 WebSocket 鎺ㄩ€佸埌 APP銆?,
        'APP HUD 瀹炴椂鏄剧ず璁粌鐘舵€侊紱鐢ㄦ埛鐐瑰嚮 Stop 鍚庯紝APP 璋冪敤 /api/v2/training/stop 缁撴潫璁粌銆?,
    ])

    doc.add_heading('涓夈€佸紑鍙戠幆澧冧笌鎶€鏈€夊瀷', level=1)
    add_table(doc, ['绫诲埆', '鎶€鏈?宸ュ叿', '璇存槑'], [
        ['鍚庣', 'Python 3銆丗astAPI銆乁vicorn銆丼QLite', 'REST API銆佽缁冪姸鎬佺鐞嗐€佹暟鎹簱涓庡疄鏃舵帴鍙?],
        ['AI 瑙嗚', 'MediaPipe Pose銆丱penCV銆丯umPy', '浜轰綋鍏抽敭鐐规娴嬨€佽搴﹁绠椼€佸姩浣滆鏁?],
        ['绉诲姩绔?, 'Kotlin銆丄ndroidX銆丆ameraX銆丷etrofit銆丱kHttp銆乄ebSocket', 'Android 鍘熺敓 APP銆佺浉鏈洪噰闆嗐€丠TTP/WS 閫氫俊'],
        ['宓屽叆寮?, 'ESP32-CAM銆丄rduino銆乄iFi銆丠TTP/MJPEG', '杈圭紭鍥惧儚閲囬泦銆佸帇缂┿€佷笂浼犮€佺姸鎬佹帶鍒?],
        ['PC 瀹㈡埛绔?, 'Python銆丱penCV銆乺equests', '鐢佃剳鎽勫儚澶寸槮瀹㈡埛绔紝璐熻矗閲囬泦鍘嬬缉鍜屼笂浼?],
        ['閮ㄧ讲娴嬭瘯', 'Windows銆丄DB銆丄ndroid Emulator銆丟radle銆丮osquitto', '鏈湴鏈嶅姟銆佹ā鎷熷櫒娴嬭瘯銆丮QTT 鏀寔'],
    ], widths=[1.2, 2.4, 3.6])

    doc.add_heading('鍥涖€佹ā鍧楄璁′笌瀹炵幇', level=1)
    doc.add_heading('4.1 鍚庣鏈嶅姟妯″潡', level=2)
    doc.add_paragraph(
        '鍚庣浠?FastAPI 涓轰富妗嗘灦锛屾牳蹇冩枃浠朵负 backend/main.py 涓?backend/main_v2_routes.py銆?
        'main.py 璐熻矗鏈嶅姟鍚姩銆佽矾鐢辨敞鍐屻€丮QTT 鍒濆鍖栧拰鍩虹鍋ュ悍妫€鏌ワ紱main_v2_routes.py 璐熻矗 v2 API锛?
        '鍖呮嫭鐢ㄦ埛璁よ瘉銆佽澶囨敞鍐屻€佽缁冭鍒掋€佽缁冩帶鍒躲€佽瑙夋帹鐞嗗拰 WebSocket 鎺ㄩ€併€?
    )
    add_table(doc, ['鎺ュ彛', '鏂规硶', '浣滅敤'], [
        ['/health', 'GET', '鍋ュ悍妫€鏌ワ紝杩斿洖鏈嶅姟鐘舵€併€佹椿璺冭澶囧拰 MQTT 杩炴帴鐘舵€?],
        ['/api/v2/auth/login', 'POST', '鐢ㄦ埛鐧诲綍锛岃繑鍥?token 涓庣敤鎴蜂俊鎭?],
        ['/api/v2/devices/register', 'POST', '娉ㄥ唽璁惧骞跺缓绔嬬敤鎴风粦瀹?],
        ['/api/v2/training/start', 'POST', '鍚姩璁粌锛岃褰?device_id銆乽ser_id銆乪xercise 涓?session_id'],
        ['/api/v2/training/stop', 'POST', '鍋滄璁粌骞舵竻鐞?active training'],
        ['/api/v2/training/active', 'GET', '鏌ョ湅褰撳墠鎵€鏈夋椿璺冭缁?],
        ['/api/v2/vision/infer/full', 'POST', '鍥惧儚鎺ㄧ悊鑱氬悎鎺ュ彛锛岃繑鍥炲叧閿偣銆佽鏁般€佽瘎鍒嗐€佸弽棣?],
        ['/ws/coach/{user_id}', 'WS', '鍚?APP 瀹炴椂鎺ㄩ€佽缁冪粨鏋?],
    ], widths=[2.4, 0.9, 3.9])
    doc.add_paragraph(
        '鏈澶氭憚鍍忓ご鎵╁睍涓紝鍚庣鍦?/infer/full 璇锋眰鍜屽搷搴斾腑澧炲姞 source銆乨evice_type銆乨evice_id 瀛楁锛?
        '浣?APP 鑳藉鍖哄垎缁撴灉鏉ユ簮锛涘悓鏃惰缁?start/stop 涓庢帹鐞嗘帴鍙ｅ繀椤讳娇鐢ㄥ悓涓€涓?device_id锛岄伩鍏嶈鏁板櫒鍜?active training 涓嶅尮閰嶃€?
    )

    doc.add_heading('4.2 AI 濮挎€佽瘑鍒笌鍔ㄤ綔璁℃暟妯″潡', level=2)
    doc.add_paragraph(
        'AI 妯″潡浣跨敤 MediaPipe Pose 杩涜浜轰綋鍏抽敭鐐规娴嬶紝鍐嶆牴鎹叧鑺傝搴﹀拰闃舵鐘舵€佸疄鐜板姩浣滆鏁般€?
        '璁℃暟閫昏緫瀵逛笉鍚屽姩浣滈厤缃笉鍚岄槇鍊硷紝渚嬪娣辫共閫氳繃鑶濆叧鑺傝搴﹀垽鏂?down/up 闃舵锛?
        '骞剁粨鍚堣繛缁抚鏈哄埗鍑忓皯璇垽銆傜郴缁熻繕璁＄畻 form_score 骞惰緭鍑轰腑鏂?feedback锛屼緵 APP 绔疄鏃跺睍绀恒€?
    )
    add_table(doc, ['杩愬姩绫诲瀷', '璇嗗埆渚濇嵁', '杈撳嚭鍐呭'], [
        ['Squat 娣辫共', '楂?鑶?韪濆叧閿偣锛岃啙鍏宠妭瑙掑害闃舵鍙樺寲', '娆℃暟銆侀樁娈点€佸Э鍔垮弽棣?],
        ['Push-up 淇崸鎾?, '鑲?鑲?鑵曞叧閿偣锛岃倶鍏宠妭瑙掑害', '娆℃暟銆佸姩浣滃畬鎴愬害銆佸Э鍔垮缓璁?],
        ['Plank 骞虫澘鏀拺', '鑲?楂?韪濊繛绾跨ǔ瀹氭€?, '淇濇寔鐘舵€併€佸Э鎬佽瘎鍒?],
        ['Lunge 寮撴', '宸﹀彸鑵胯搴︺€佽韩浣撻噸蹇?, '娆℃暟涓庡乏鍙充晶鍔ㄤ綔璐ㄩ噺'],
        ['Curl/Press 绛夊姏閲忓姩浣?, '鑲樸€佽偐鍏宠妭瑙掑害鍙樺寲', '娆℃暟銆侀樁娈靛拰鍙嶉'],
    ], widths=[1.5, 3.1, 2.6])

    doc.add_heading('4.3 Android APP 妯″潡', level=2)
    doc.add_paragraph(
        'Android APP 浣跨敤 Kotlin 寮€鍙戯紝鏍稿績璁粌椤典负 TrainingFragment銆傝缁冮〉鎻愪緵鎽勫儚澶存潵婧愰€夋嫨銆佽繍鍔ㄧ被鍨嬮€夋嫨銆?
        '寮€濮?鍋滄璁粌鎸夐挳銆侀鏋堕瑙堛€佽鏁般€佽瘎鍒嗗拰鏁欑粌鍙嶉銆傜綉缁滃眰浣跨敤 Retrofit/OkHttp锛屽疄鏃舵帹閫佷娇鐢?WebSocket銆?
    )
    add_table(doc, ['鏂囦欢/缁勪欢', '鍔熻兘'], [
        ['TrainingFragment.kt', '璁粌椤典富閫昏緫锛氭憚鍍忓ご鏉ユ簮鍒囨崲銆佽缁?start/stop銆佹帹鐞嗚姹傘€丠UD 鏇存柊'],
        ['fragment_training.xml', '璁粌椤靛竷灞€锛氶瑙堝尯鍩熴€乻ource spinner銆乪xercise spinner銆佽缁冩寜閽€丠UD'],
        ['CameraCapture.kt', '鍩轰簬 CameraX 鐨勬墜鏈烘湰鏈烘憚鍍忓ご閲囬泦銆乺esize銆丣PEG 鍘嬬缉鍜?Base64 缂栫爜'],
        ['MjpegClient.kt', 'ESP32-CAM MJPEG 娴佽鍙栦笌鎶藉抚'],
        ['ApiClient.kt / ApiService.kt', 'HTTP 鎺ュ彛灏佽銆乀oken 鍜?base_url 绠＄悊'],
        ['Models.kt', '璇锋眰/鍝嶅簲鏁版嵁妯″瀷锛屽寘鎷?VisionInferRequest銆乀rainingStartRequest 绛?],
    ], widths=[2.3, 4.9])
    doc.add_paragraph('澶氭憚鍍忓ご閫夋嫨鍣ㄥ寘鍚笁涓€夐」锛欵SP32-CAM銆丳hone Camera銆丳C Camera銆傚垏鎹㈡潵婧愭椂 APP 浼氶噴鏀句笂涓€涓潵婧愮殑璧勬簮锛屽苟娓呯┖鏈湴涓存椂璁℃暟锛岄槻姝笉鍚屾潵婧愪箣闂寸殑璁粌鐘舵€佷覆鎵般€?)

    doc.add_heading('4.4 ESP32-CAM 妯″潡', level=2)
    doc.add_paragraph(
        'ESP32-CAM 妯″潡璐熻矗纭欢绔浘鍍忛噰闆嗗拰杈圭紭鍘嬬缉銆傜郴缁熶繚鐣欏師鏈?MJPEG 娴佽矾寰勶紝APP 鍙粠 '
        'http://<esp32_ip>:81/stream 鎷夊彇棰勮锛屽悓鏃?ESP32 涔熷彲鎸夎缁冪姸鎬佽皟鏁翠笂浼犻鐜囥€?
        '鍥轰欢鏀寔 WiFi 杩炴帴銆丠TTP 涓婁紶銆佸姩鎬?JPEG 璐ㄩ噺銆佸抚澶у皬瀹堝崼銆佹柇绾块噸杩炲拰 OTA銆?
    )
    add_bullets(doc, [
        '榛樿 device_id锛歟sp32cam-001銆?,
        '鍥惧儚鍘嬬缉锛氬湪 ESP32 绔畬鎴?JPEG 鍘嬬缉锛岄檷浣庝笂浼犲甫瀹姐€?,
        '鏈嶅姟绔帶鍒讹細鍚庣鍙€氳繃 next_interval_ms 鍛婄煡璁惧鍦ㄨ缁冧腑/鏈缁冩椂璋冩暣閲囨牱棰戠巼銆?,
        '婕旂ず鎻愰啋锛氭崲缃戠粶鍚庨渶瑕侀噸鏂扮‘璁?Backend URL銆丒SP32 IP銆乨evice_id 涓?APK 涓嬭浇鍦板潃銆?,
    ])

    doc.add_heading('4.5 澶氭憚鍍忓ご鏉ユ簮鎵╁睍', level=2)
    doc.add_paragraph(
        '鏈€缁堢増鏈皢鎽勫儚澶存潵婧愮粺涓€鎶借薄涓?source + device_id銆傛棤璁哄浘鍍忔潵鑷?ESP32銆佹墜鏈鸿繕鏄?PC锛?
        '鏈€缁堝潎璋冪敤鍚屼竴涓悗绔帴鍙?/api/v2/vision/infer/full锛屽苟澶嶇敤鍚屼竴濂楄缁冪姸鎬併€佽鏁板櫒鍜?WebSocket 鎺ㄩ€侀€昏緫銆?
    )
    add_table(doc, ['鏉ユ簮', '閲囬泦绔?, '鍘嬬缉鍙傛暟', 'device_id', 'APP 琛屼负'], [
        ['ESP32-CAM', 'ESP32-CAM OV2640', 'ESP32 绔?JPEG锛孉PP 鎷?MJPEG 鍚庡彲鎶藉抚', 'esp32cam-001', '鏄剧ず ESP32 棰勮锛屾娊甯т笂浼?鎺ユ敹鎺ㄩ€?],
        ['Phone Camera', 'Android CameraX', 'maxWidth=640锛孞PEG quality=58锛宨nterval=500ms', 'phone-<android-device-id>', '鏄剧ず鏈満 CameraX 棰勮骞剁洿鎺ヤ笂浼?],
        ['PC Camera', 'Python OpenCV', '榛樿 maxWidth=640锛孞PEG quality=60锛宨nterval=500ms', 'pc-camera-001', 'APP 涓嶉噰闆嗭紝鍙闃呭悗绔?HUD锛汸C agent 璐熻矗涓婁紶'],
    ], widths=[1.3, 1.6, 2.2, 1.6, 2.3])

    doc.add_paragraph('缁熶竴璇锋眰绀轰緥锛?)
    add_code(doc, r'''
POST /api/v2/vision/infer/full
{
  "image": "base64 jpeg",
  "device_id": "esp32cam-001 | phone-xxx | pc-camera-001",
  "exercise": "squat",
  "source": "esp32cam | phone | pc",
  "backend": "mediapipe"
}
''')

    doc.add_heading('4.6 PC Camera 鐦﹀鎴风', level=2)
    doc.add_paragraph(
        'PC Camera 閫氳繃 pc_simulator/pc_camera_agent.py 瀹炵幇銆傝鑴氭湰浣跨敤 OpenCV 鎵撳紑鏈満鎽勫儚澶达紝'
        '鎸夋寚瀹氶棿闅旇鍙栫敾闈€佺缉鏀俱€丣PEG 缂栫爜銆丅ase64 鍖栵紝鐒跺悗璋冪敤鍚庣缁熶竴鎺ㄧ悊鎺ュ彛銆?
        'APP 閫夋嫨 PC Camera 鍚庝笉鍐嶉噰闆嗗浘鍍忥紝鑰屾槸閫氳繃 WebSocket 鎺ユ敹 pc-camera-001 鐨勮缁冪粨鏋溿€?
    )
    add_code(doc, r'''
cd C:\Users\hjl\.openclaw\workspace\smart_fitness
python pc_simulator\pc_camera_agent.py --server http://127.0.0.1:8080 --device-id pc-camera-001 --exercise squat --preview
''')

    doc.add_heading('浜斻€佸叧閿疄鐜伴棶棰樹笌瑙ｅ喅鏂规', level=1)
    add_table(doc, ['闂', '鍘熷洜', '瑙ｅ喅鏂规'], [
        ['璁粌寮€濮嬭繑鍥?400', '鍚庣 training/start 瑕佹眰 device_id + user_id锛孉PP 鏃ц姹傛病鏈変紶 user_id', 'TrainingStartRequest 澧炲姞 userId/source锛宻tartTraining 浣跨敤 ApiClient.userId'],
        ['Phone Camera 棰勮鍙兘閬尅鎸夐挳', 'PreviewView 榛樿 SurfaceView 鍙兘浣嶄簬瑙︽懜灞備笂鏂?, 'fragment_training.xml 涓缃?implementationMode="compatible"'],
        ['涓嶅悓鎽勫儚澶磋鏁颁覆鎵?, '澶氫釜鏉ユ簮鍏辩敤鏃?device_id 鎴栨湰鍦拌鏁扮姸鎬?, '鎸夋潵婧愮敓鎴愮嫭绔?device_id锛屽垏鎹?source 鏃堕噸缃湰鍦?HUD'],
        ['鍚庣閲嶅惎鍚?WS 鍋囧湪绾?, 'APP 鍙垽鏂?WebSocket 瀵硅薄闈炵┖锛屾湭鍒ゆ柇鐪熷疄杩炴帴鐘舵€?, '璁粌寮€濮嬫椂寮哄埗閲嶈繛鎴栨娴嬭繛鎺ユ湁鏁堟€?],
        ['鏃犱汉浣撳抚瀵艰嚧璁℃暟褰掗浂', 'detected=false 鏃跺搷搴?rep_count 鍙兘涓?0锛孉PP 鐩存帴瑕嗙洊 UI', 'APP 瀵?no-person 甯т繚鐣?last valid reps'],
        ['PC 鎽勫儚澶翠笌 APP 浼氳瘽涓嶄竴鑷?, 'PC agent device_id 涓?APP PC device_id 涓嶄竴鑷?, '璁剧疆椤甸厤缃?pc_device_id锛岃剼鏈?--device-id 蹇呴』涓€鑷?],
    ], widths=[1.7, 2.3, 3.2])

    doc.add_heading('鍏€佹祴璇曚笌楠岃瘉', level=1)
    doc.add_heading('6.1 鏋勫缓涓庨潤鎬佹鏌?, level=2)
    add_table(doc, ['娴嬭瘯椤?, '鍛戒护/鏂瑰紡', '缁撴灉'], [
        ['Python 璇硶妫€鏌?, 'python -m py_compile backend\\main_v2_routes.py pc_simulator\\pc_camera_agent.py', '閫氳繃'],
        ['Android 鏋勫缓', '.\\gradlew.bat :app:assembleDebug', '閫氳繃锛岀敓鎴?debug APK'],
        ['APK 澶嶅埗', '澶嶅埗涓?smart_fitness_pose_preview.apk 涓?smart_fitness_multicamera.apk', '閫氳繃'],
        ['鍚庣鍋ュ悍妫€鏌?, 'GET /health', '杩斿洖 healthy锛孧QTT connected'],
    ], widths=[1.6, 3.8, 1.8])

    doc.add_heading('6.2 妯℃嫙鍣ㄧ鍔熻兘娴嬭瘯', level=2)
    add_table(doc, ['鍦烘櫙', '楠岃瘉杩囩▼', '缁撴灉'], [
        ['Training 椤?UI', 'ADB + uiautomator 妫€鏌?esp32_preview銆乻pinner_camera_source銆乻pinner_exercise銆乥tn_training_toggle', '鎺т欢瀛樺湪'],
        ['鏉ユ簮涓嬫媺妗?, '鐐瑰嚮 Camera Source spinner', '鏄剧ず ESP32-CAM銆丳hone Camera銆丳C Camera 涓夐」'],
        ['Phone Camera 鎺ㄧ悊', '鍒囧埌 Phone Camera锛孋ameraX 閲囬泦鍚庢寔缁?POST /infer/full', '鍚庣 200 OK锛岀害 500ms/甯?],
        ['Phone Camera 璁粌', '鐐瑰嚮 Start锛岃皟鐢?/training/start锛岃瀵?active training 涓?UI 鐘舵€?, '鎸夐挳鍙?Stop锛岃鏃跺櫒杩愯锛宎ctive 鏈夎褰?],
        ['Phone Camera 鍋滄', '鐐瑰嚮 Stop锛岃皟鐢?/training/stop', 'active 娓呯┖锛岃缁冪粨鏉?],
        ['PC Camera 妯″紡', 'APP 閫夋嫨 PC Camera锛屽惎鍔?pc_camera_agent.py 涓婁紶鍥惧儚', '鍚庣鏀跺埌 pc-camera-001 鎺ㄧ悊锛孉PP HUD 鏄剧ず score/feedback'],
        ['宕╂簝妫€鏌?, 'logcat 杩囨护 FATAL EXCEPTION/AndroidRuntime', '鏈彂鐜板穿婧?],
    ], widths=[1.6, 3.8, 1.8])

    doc.add_heading('6.3 澶氭憚鍍忓ご闂幆缁撴灉', level=2)
    add_table(doc, ['鏉ユ簮', '棰勮/閲囬泦', '鎺ㄧ悊涓婁紶', '璁粌鎺у埗', '缁撹'], [
        ['ESP32-CAM', 'MJPEG 璺緞淇濈暀', '鍘熻矾寰勫吋瀹?, 'device_id=esp32cam-001', '寰呯‖浠剁幇鍦哄洖褰掞紝浠ｇ爜鏈牬鍧?],
        ['Phone Camera', 'CameraX 棰勮姝ｅ父', 'POST /infer/full 200 OK', 'Start/Stop 姝ｅ父锛孒UD 鏇存柊', '閫氳繃'],
        ['PC Camera', 'OpenCV 閲囬泦姝ｅ父', 'PC agent 涓婁紶 200 OK', 'pc-camera-001 active training 姝ｅ父锛孒UD 鏀跺埌缁撴灉', '閫氳繃'],
    ], widths=[1.3, 1.8, 1.8, 2.2, 1.1])

    doc.add_heading('涓冦€侀儴缃蹭笌浣跨敤璇存槑', level=1)
    doc.add_heading('7.1 鍚姩鍚庣鏈嶅姟', level=2)
    add_code(doc, r'''
cd C:\Users\hjl\.openclaw\workspace\smart_fitness\backend
python -m uvicorn main:app --host 0.0.0.0 --port 8080
''')
    doc.add_heading('7.2 鍚姩 MQTT Broker', level=2)
    add_code(doc, r'''
cd "C:\Program Files\mosquitto"
.\mosquitto.exe -v
''')
    doc.add_heading('7.3 鍚姩 APK 闈欐€佷笅杞芥湇鍔?, level=2)
    add_code(doc, r'''
cd C:\Users\hjl\.openclaw\workspace\smart_fitness
python -m http.server 8090
''')
    doc.add_heading('7.4 瀹夎 APK', level=2)
    add_code(doc, r'''
C:\Users\hjl\AppData\Local\Android\Sdk\platform-tools\adb.exe install -r C:\Users\hjl\.openclaw\workspace\smart_fitness\smart_fitness_pose_preview.apk
''')
    doc.add_heading('7.5 鎹㈢綉缁滈厤缃彁閱?, level=2)
    add_bullets(doc, [
        'Backend URL锛氬～鍐?PC 褰撳墠灞€鍩熺綉 IP + :8080锛屼緥濡?http://192.168.x.x:8080/銆?,
        'ESP32 IP锛氬～鍐?ESP32-CAM 鍦ㄥ悓涓€缃戠粶涓嬬殑鏂?IP銆?,
        'ESP32 device_id锛氶€氬父淇濇寔 esp32cam-001銆?,
        'APK/static 鏈嶅姟锛氶€氬父涓?PC 褰撳墠灞€鍩熺綉 IP + :8090銆?,
        'MQTT锛氶€氬父浠嶄负 1883 绔彛銆?,
    ])

    doc.add_heading('鍏€侀」鐩垚鏋滀笌浜や粯鐗?, level=1)
    add_table(doc, ['浜や粯鐗?, '璺緞/鍚嶇О', '璇存槑'], [
        ['Android APK', 'smart_fitness_pose_preview.apk / smart_fitness_multicamera.apk', '宸叉瀯寤虹殑 Android 瀹夎鍖?],
        ['鍚庣鏈嶅姟', 'backend/main.py銆乥ackend/main_v2_routes.py', 'FastAPI 鏈嶅姟涓?v2 API'],
        ['Android 瀹㈡埛绔簮鐮?, 'android_app/app/src/main/...', 'Kotlin APP 婧愮爜'],
        ['PC 鎽勫儚澶村鎴风', 'pc_simulator/pc_camera_agent.py', 'OpenCV 鎽勫儚澶翠笂浼?agent'],
        ['ESP32-CAM 鍥轰欢', 'edge_esp32cam/esp32cam_fitness/esp32cam_fitness.ino', '宓屽叆寮忔憚鍍忓ご鍥轰欢'],
        ['澶氭憚鍍忓ご璇存槑', 'MULTI_CAMERA_SOURCES.md', '澶氭潵婧愭帴鍏ヤ笌浣跨敤璇存槑'],
        ['瀹炶鎶ュ憡', '鏅鸿兘鍋ヨ韩鎸囧绯荤粺瀹炶鎶ュ憡_鏈€缁堢増.docx', '鏈姤鍛?Word 鐗堟湰'],
    ], widths=[1.6, 3.2, 2.4])

    doc.add_heading('涔濄€佸疄璁€荤粨', level=1)
    doc.add_paragraph(
        '鏈瀹炶瀹屾垚鐨勪笉鍙槸鍗曠偣鍔熻兘锛岃€屾槸涓€涓法绔棴鐜郴缁燂細浠庡浘鍍忛噰闆嗐€佺綉缁滀紶杈撱€佸悗绔帹鐞嗐€佸姩浣滆鏁般€?
        '瀹炴椂鎺ㄩ€佸埌绉诲姩绔睍绀哄潎宸叉墦閫氥€傜郴缁熷湪璁捐涓婇€夋嫨浜嗏€滅渚у帇缂?+ 鍚庣缁熶竴鎺ㄧ悊鈥濈殑鏂规锛?
        '鏃㈤檷浣庝簡甯﹀鍘嬪姏锛屼篃閬垮厤鍦ㄦ瘡涓噰闆嗙閲嶅閮ㄧ讲濮挎€佹ā鍨嬶紝渚夸簬鎵╁睍鏂版憚鍍忓ご鏉ユ簮銆?
    )
    doc.add_paragraph(
        '鍦ㄨ皟璇曡繃绋嬩腑锛屼富瑕侀毦鐐归泦涓湪 device_id 涓€鑷存€с€佽缁冪姸鎬佸悓姝ャ€乄ebSocket 瀹炴椂鎺ㄩ€併€丄ndroid 鐩告満璧勬簮閲婃斁銆?
        '浠ュ強鏃犱汉浣撳抚瀵艰嚧 UI 鐘舵€佽鏇存柊绛夐棶棰樸€傞€氳繃灏嗘憚鍍忓ご鏉ユ簮鎶借薄涓?source/device_id锛?
        '骞惰姹?training/start銆乼raining/stop 涓?/infer/full 浣跨敤鍚屼竴涓?device_id锛岀郴缁熸渶缁堝舰鎴愪簡绋冲畾鐨勫鎽勫儚澶磋缁冮棴鐜€?
    )
    doc.add_paragraph(
        '鍚庣画鍙互缁х画瀹屽杽鐪熷疄 ESP32-CAM 鐜板満鍥炲綊銆佹紨绀鸿棰戝綍鍒躲€佹洿澶氬姩浣滆鍒欍€佽缁冨巻鍙茬粺璁″浘琛ㄥ拰鑷姩鍖栨祴璇曡剼鏈€?
        '鏁翠綋鏉ョ湅锛屾湰椤圭洰宸茬粡婊¤冻瀹炶瀵光€滅‖浠舵劅鐭ャ€丄I 璇嗗埆銆佺Щ鍔ㄧ浜や簰銆佸悗绔湇鍔′笌绯荤粺闆嗘垚鈥濈殑缁煎悎瑕佹眰銆?
    )

    doc.add_heading('闄勫綍 A锛氬叧閿矾寰勬竻鍗?, level=1)
    add_code(doc, r'''
C:\Users\hjl\.openclaw\workspace\smart_fitness\backend\main.py
C:\Users\hjl\.openclaw\workspace\smart_fitness\backend\main_v2_routes.py
C:\Users\hjl\.openclaw\workspace\smart_fitness\android_app\app\src\main\java\com\smartfitness\app\ui\training\TrainingFragment.kt
C:\Users\hjl\.openclaw\workspace\smart_fitness\android_app\app\src\main\res\layout\fragment_training.xml
C:\Users\hjl\.openclaw\workspace\smart_fitness\pc_simulator\pc_camera_agent.py
C:\Users\hjl\.openclaw\workspace\smart_fitness\edge_esp32cam\esp32cam_fitness\esp32cam_fitness.ino
C:\Users\hjl\.openclaw\workspace\smart_fitness\MULTI_CAMERA_SOURCES.md
''')

    doc.add_heading('闄勫綍 B锛氬綋鍓嶇増鏈檺鍒?, level=1)
    add_bullets(doc, [
        'ESP32-CAM 瀹屾暣鐪熸満鍥炲綊渚濊禆纭欢涓婄數鍜屽悓涓€灞€鍩熺綉鐜銆?,
        'PC Camera 渚濊禆鏈満鎽勫儚澶村彲鐢紝鑻ユ憚鍍忓ご琚叾浠栫▼搴忓崰鐢紝闇€瑕侀噴鏀惧悗閲嶅惎 agent銆?,
        '涓嶅悓缃戠粶鐜涓嬮渶瑕侀噸鏂伴厤缃?Backend URL銆丒SP32 IP 涓?APK 涓嬭浇鍦板潃銆?,
        '褰撳墠鍔ㄤ綔璁℃暟瑙勫垯浠ヨ搴﹂槇鍊煎拰闃舵鐘舵€佷负涓伙紝澶嶆潅鍔ㄤ綔鍙户缁姞鍏ユ洿绮剧粏鐨勬椂搴忔ā鍨嬨€?,
    ])

    doc.save(OUT)
    print(str(OUT))


if __name__ == '__main__':
    main()
